from .constant import COVER_CHAR, STEG_CHAR

import cv2
import ffmpeg
import numpy as np
import openpyxl
import os
import re
import wave

from docx import Document
from docx.shared import RGBColor
from openpyxl.styles import Font
from PIL import Image, ImageSequence
from PyPDF2 import PdfReader, PdfWriter


DELIMITER = '====='
DELIMITER_IDX = -len(DELIMITER)


def to_bin(data):
    """Convert `data` to binary format as string"""
    if isinstance(data, str):
        return ''.join([format(ord(i), "08b") for i in data])
    elif isinstance(data, bytes) or isinstance(data, np.ndarray):
        return [format(i, "08b") for i in data]
    elif isinstance(data, int) or isinstance(data, np.uint8):
        return format(data, "08b")
    else:
        raise TypeError("Type not supported.")


def encode_gif(gif_path, output_path, secret_data, lsb=1):

    mask = 255 - pow(2, lsb) + 1
    # open GIF file and copy frames into array
    gif = Image.open(gif_path)

    frames = [frame.copy() for frame in ImageSequence.Iterator(gif)]

    # get GIF properties
    n_frames = gif.n_frames
    height = gif.height
    width = gif.width

    usable = 0
    for frame in frames:
        if frame.mode == 'RGB' or 'RGBA':
            usable += 1
    # maximum bytes to encode
    max_encode_bytes = usable * height * width * 3 // 8

    print("[*] Maximum bytes to encode:", max_encode_bytes)

    # append delimiter
    secret_data += DELIMITER
    secret_data = secret_data.encode("utf-8")

    # check if payload is larger than maximum bytes
    if len(secret_data) > max_encode_bytes:
        raise ValueError(
            "[!] Insufficient bytes, need bigger gif file or less data.")

    print("[*] Encoding data...")

    binary_secret_data = "".join(to_bin(secret_data))
    binary_secret_data += (lsb - (len(binary_secret_data) % lsb)) * '0'

    data_index = 0
    frame_counter = 0
    x = 0
    y = 0

    secret_data_list = [int(binary_secret_data[i:i+lsb], 2)
                        for i in range(0, len(binary_secret_data), lsb)]
    data_len = len(secret_data_list)

    # encode first pixel of every frame, then move to next bit of every frame
    while data_index < data_len:
        # ignore P frames
        if frames[frame_counter].mode in ['RGBA', 'RGB']:
            
            if frames[frame_counter].mode == 'RGBA':
                frames[frame_counter] = frames[frame_counter].convert('RGB')

            # pixel data
            r, g, b = frames[frame_counter].getpixel((x, y))
            
            if data_index < data_len:
                # least significant red pixel bit(s)
                r = (r & mask) | secret_data_list[data_index]
                data_index += 1
            if data_index < data_len:
                # least significant green pixel bit(s)
                g = (g & mask) | secret_data_list[data_index]
                data_index += 1
            if data_index < data_len:
                # least significant blue pixel bit(s)
                b = (b & mask) | secret_data_list[data_index]
                data_index += 1            

            # update pixel data in frame
            frames[frame_counter].putpixel((x, y), (r, g, b))

        frame_counter += 1
        if frame_counter == n_frames:
            frame_counter = 0
            x += 1
            if x == width: 
                x = 0
                y += 1


    # save encoded gif
    if n_frames > 1:
        frames[0].save(output_path, save_all=True,
                       append_images=frames[1:])
    else:
        frames[0].save(output_path)


def decode_gif(gif_path, lsb=1):
    print("[+] Decoding...")

    # open GIF file and copy frames into array
    gif = Image.open(gif_path)
    frames = [frame.copy() for frame in ImageSequence.Iterator(gif)]

    # get GIF properties
    n_frames = gif.n_frames
    height = gif.height
    width = gif.width

    binary_data = ""
    decoded_data = ""

    buffer = ""
    byte_count = 1
    continuation_byte_count = 0
    
    frame_counter = 0
    x = 0
    y = 0
    while y != height and x != width and frame_counter != n_frames:
        # ignore P frames
        if frames[frame_counter].mode == 'RGB':

            # pixel data
            r, g, b = frames[frame_counter].getpixel((x, y))

            # append lsb(s) into binary data array
            binary_data += to_bin(r)[-lsb:]
            binary_data += to_bin(g)[-lsb:]
            binary_data += to_bin(b)[-lsb:]

            # decode binary data if more than 8
            while len(binary_data) >= 8:
                # slice the first 8 bits away 
                payload_byte = binary_data[:8]
                binary_data = binary_data[8:]
                
                if(payload_byte[0] == '0'):
                    byte_count = 1
                    # append character to decoded
                    decoded_data += chr(int(payload_byte, 2))
                elif(payload_byte[:3] == '110'):
                    byte_count = 2
                    buffer += payload_byte
                elif(payload_byte[:4] == '1110'):
                    byte_count = 3
                    buffer += payload_byte
                elif(payload_byte[:5] == '11110'):
                    byte_count = 4
                    buffer += payload_byte
                elif(payload_byte[:2] == '10'):
                    buffer += payload_byte
                    continuation_byte_count += 1
                    if(continuation_byte_count == byte_count - 1):
                        # append character to decoded
                        decoded_data += bytes(int(buffer[i:i+8], 2) for i in range(0, len(buffer), 8)).decode('utf-8')
                        buffer = ""
                        byte_count = 1
                        continuation_byte_count = 0

                if decoded_data[DELIMITER_IDX:] == DELIMITER:
                    return decoded_data[:DELIMITER_IDX]

        frame_counter += 1
        if frame_counter == n_frames:
            frame_counter = 0
            x += 1
            if x == width:
                x = 0
                y += 1
    raise ValueError(
        "Could not decode. (Could not find the delimiter '%s')" % DELIMITER)


def encode_image(image_name, secret_data, lsb=1):

    # create mask
    mask = 255 - pow(2, lsb) + 1

    image = cv2.imread(image_name)
    # maximum bytes to encode
    n_bytes = image.shape[0] * image.shape[1] * 3 // 8
    print("[*] Maximum bytes to encode:", n_bytes)

    # add stopping criteria
    secret_data += DELIMITER
    secret_data = secret_data.encode("utf-8")
    data_index = 0

    if len(secret_data) > n_bytes:
        raise ValueError(
            "[!] Insufficient bytes, need bigger image or less data.")
    print("[*] Encoding data...")

    # convert data to binary
    binary_secret_data = "".join(to_bin(secret_data))
    pad = (lsb - (len(binary_secret_data) % lsb)) * '0'
    binary_secret_data += pad
    # convert secret data to int array
    secret_data_list = [int(binary_secret_data[i:i+lsb], 2)
                        for i in range(0, len(binary_secret_data), lsb)]
    # size of data to hide

    data_len = len(secret_data_list)
    for row in image:
        for pixel in row:
            # convert RGB values to binary format
            r, g, b = pixel
            # modify the least significant bit only if there is still data to store
            if data_index < len(secret_data_list):
                # least significant red pixel bit
                pixel[0] = (r & mask) | secret_data_list[data_index]
                data_index += 1
            if data_index < data_len:
                # least significant green pixel bit
                pixel[1] = (g & mask) | secret_data_list[data_index]
                data_index += 1
            if data_index < data_len:
                # least significant blue pixel bit
                pixel[2] = (b & mask) | secret_data_list[data_index]
                data_index += 1
            # if data is encoded, just break out of the loop
            if data_index >= data_len:
                break
    return image


def decode_image(image_name, lsb=1):
    print("[+] Decoding...")
    # read the image
    image = cv2.imread(image_name)
    binary_data = ""
    decoded_data = ""

    buffer = ""
    byte_count = 1
    continuation_byte_count = 0

    for row in image:
        for pixel in row:
            r, g, b = to_bin(pixel)
            binary_data += r[-lsb:]
            binary_data += g[-lsb:]
            binary_data += b[-lsb:]

            # decode binary data if more than 8
            while len(binary_data) >= 8:
                # slice the first 8 bits away 
                payload_byte = binary_data[:8]
                binary_data = binary_data[8:]
                
                if(payload_byte[0] == '0'):
                    byte_count = 1
                    # append character to decoded
                    decoded_data += chr(int(payload_byte, 2))
                elif(payload_byte[:3] == '110'):
                    byte_count = 2
                    buffer += payload_byte
                elif(payload_byte[:4] == '1110'):
                    byte_count = 3
                    buffer += payload_byte
                elif(payload_byte[:5] == '11110'):
                    byte_count = 4
                    buffer += payload_byte
                elif(payload_byte[:2] == '10'):
                    
                    buffer += payload_byte
                    continuation_byte_count += 1
                    if(continuation_byte_count == byte_count - 1):
                        # append character to decoded
                        decoded_data += bytes(int(buffer[i:i+8], 2) for i in range(0, len(buffer), 8)).decode('utf-8')
                        buffer = ""
                        byte_count = 1
                        continuation_byte_count = 0

                if decoded_data[DELIMITER_IDX:] == DELIMITER:
                    return decoded_data[:DELIMITER_IDX]
    raise ValueError(
        "Could not decode. (Could not find the delimiter '%s')" % DELIMITER)



def encode_audio(audio_path, output_path, secret_data, lsb=1):

    # create mask
    mask = 255 - pow(2, lsb) + 1

    # Open the input wave file
    with wave.open(audio_path, 'rb') as song:
        # Read frames and convert to byte array
        frame_bytes = bytearray(list(song.readframes(song.getnframes())))
        n_frames = len(frame_bytes)

        # Append delimiter
        secret_data += DELIMITER
        secret_data = secret_data.encode("utf-8")

        # Convert text to bit array
        bits = []
        for byte in secret_data:
            for i in range(7, -1, -1):  # Iterate from the MSB to the LSB
                bit = (byte >> i) & 0x01
                bits.append(bit)
        n_bits = len(bits)

        # add padding bits !important
        while n_bits % lsb != 0:
            bits.append(0)
            n_bits += 1

        # get bits
        joined = [int(binary, 2) for binary in [''.join(
            map(str, bits[i:i+lsb])) for i in range(0, len(bits), lsb)]]

        # we want to modify each LSB of each byte of the data (frame), so maximum bits is number of frame bytes
        print(f"[*] Encoding {n_bits}/{n_frames} bits of data in audio.")
        if n_bits > n_frames:
            raise ValueError(
                "[!] Insufficient bytes, need bigger image or less data.")

        # Replace LSB of each byte of the audio data by one bit from the text bit array
        for i, bit in enumerate(joined):

            frame_bytes[i] = (frame_bytes[i] & mask) | bit
        # Get the modified bytes
        frame_modified = bytes(frame_bytes)

        # Write bytes to a new wave audio file
        with wave.open(output_path, 'wb') as fd:
            fd.setparams(song.getparams())
            fd.writeframes(frame_modified)
        song.close()


def decode_audio(audio_path, lsb=1):
    song = wave.open(audio_path, mode='rb')

    mask = 0
    i = 0
    while i != lsb:
        mask += pow(2, i)
        i += 1
    # Convert audio to byte array
    frame_bytes = bytearray(list(song.readframes(song.getnframes())))

    decoded_data = ""
    # Extract the LSB of each byte
    extracted = [frame_bytes[i] & mask for i in range(len(frame_bytes))]

    extracted = [format(num, '0' + str(lsb) + 'b') for num in extracted]

    # for each byte, convert it to a char and then add it to the final string
    # we convert it to a char by joining 8 bits together and converting it to an int, then obtaining its char

    binary_data = ""
    song.close()

    buffer = ""
    byte_count = 1
    continuation_byte_count = 0
    for i in extracted:
        binary_data += i
        while len(binary_data) >= 8:
            # slice the first 8 bits away 
            payload_byte = binary_data[:8]
            binary_data = binary_data[8:]
            
            
            if(payload_byte[0] == '0'):
                byte_count = 1
                # append character to decoded
                decoded_data += chr(int(payload_byte, 2))
            elif(payload_byte[:3] == '110'):
                byte_count = 2
                buffer += payload_byte
            elif(payload_byte[:4] == '1110'):
                byte_count = 3
                buffer += payload_byte
            elif(payload_byte[:5] == '11110'):
                byte_count = 4
                buffer += payload_byte
            elif(payload_byte[:2] == '10'):
                buffer += payload_byte
                continuation_byte_count += 1
                if(continuation_byte_count == byte_count - 1):
                    # append character to decoded
                    decoded_data += bytes(int(buffer[i:i+8], 2) for i in range(0, len(buffer), 8)).decode('utf-8')
                    buffer = ""
                    byte_count = 1
                    continuation_byte_count = 0
            if decoded_data.endswith(DELIMITER):
                return decoded_data[:DELIMITER_IDX]

    raise ValueError(
        "Could not decode. (Could not find the delimiter '%s')" % DELIMITER)


def encode_text(text_path, output_path, secret_data, lsb=1):
    # Get contents of cover file
    cover_file = open(text_path, 'r', encoding="utf-8")
    cover_text = cover_file.read()
    cover_file.close()

    # Get maximum bytes to encode (only encode before and after selected cover character)
    max_encode_bytes = cover_text.count(COVER_CHAR) * 2 * lsb // 8
    print("[*] Maximum bytes to encode:", max_encode_bytes)

    # Append delimiter to payload
    secret_data += DELIMITER
    bytes_secret_data = bytes(secret_data, 'utf-8')

    # Check if payload is larger than maximum bytes
    if len(bytes_secret_data) > max_encode_bytes:
        raise ValueError(
            "[!] Insufficient bytes, need bigger text file or less data.")

    # Convert payload to binary string
    print("[*] Encoding data...")
    binary_secret_data = ''.join(i for i in to_bin(bytes_secret_data))

    # Initialise sentinel values
    len_payload = len(binary_secret_data)
    idx_payload = 0
    steg_text = ''
    base_steg_char = ''.join(i for i in to_bin(
        bytes(STEG_CHAR, 'utf-8')))[:-lsb]

    # Loop through each character of cover text
    for char in cover_text:
        # Check if end of payload is reached
        if (idx_payload >= len_payload):
            steg_text += char
        else:
            # Check for cover character
            if (char == COVER_CHAR):
                # Get next set of bits to encode (padded with 0 if not enough bits)
                encode = binary_secret_data[idx_payload:idx_payload +
                                            lsb].ljust(lsb, '0')
                # Get payload character
                payload_char = bytes.fromhex(
                    hex(int(base_steg_char + encode, 2))[2:]).decode('utf-8')
                # Add payload character to stego text
                steg_text += payload_char
                idx_payload += lsb
                steg_text += char

                # Check if end of payload is reached before encoding again (two encodings are performed for each cover character, before and after it)
                if (idx_payload < len_payload):
                    # Get next set of bits to encode (padded with 0 if not enough bits)
                    encode = binary_secret_data[idx_payload:idx_payload +
                                                lsb].ljust(lsb, '0')
                    # Get payload character
                    payload_char = bytes.fromhex(
                        hex(int(base_steg_char + encode, 2))[2:]).decode('utf-8')
                    # Add payload character to stego text
                    steg_text += payload_char
                    idx_payload += lsb
            else:
                steg_text += char

    # Write stego text to output file
    with open(output_path, 'w', encoding="utf-8") as f:
        f.write(steg_text)


def decode_text(text_path, lsb=1):
    # Get contents of stego file
    steg_file = open(text_path, 'r', encoding="utf-8")
    steg_text = steg_file.read()
    steg_file.close()
    print("[+] Decoding...")

    # Initialise sentinel values
    esc_cover_char = re.escape(COVER_CHAR)
    regex = r"[^{}]{}[^{}]".format(
        esc_cover_char, esc_cover_char, esc_cover_char)
    base_steg_char = ''.join(i for i in to_bin(
        bytes(STEG_CHAR, 'utf-8')))[:-lsb]
    binary_payload = ""

    # Get payload characters
    list_payload_char = re.findall(regex, steg_text)
    for raw in list_payload_char:
        list_raw = raw.split(COVER_CHAR)
        for char in list_raw:
            # Convert char to binary
            char = ''.join(i for i in to_bin(bytes(char, 'utf-8')))
            # Populate list of steg characters
            if (char[:-lsb] == base_steg_char):
                binary_payload += char[-lsb:]

    # Convert payload to list of bytes
    list_payload_bytes = [binary_payload[i:i + 8]
                          for i in range(0, len(binary_payload), 8)]

    # Convert payload to characters
    decoded_payload = ""
    buffer = ""
    byte_count = 1
    continuation_byte_count = 0
    for payload_byte in list_payload_bytes:
        # Check variable byte encoding
        if(payload_byte[0] == '0'):
            byte_count = 1
            decoded_payload += chr(int(payload_byte, 2))
        elif(payload_byte[:3] == '110'):
            byte_count = 2
            buffer += payload_byte
        elif(payload_byte[:4] == '1110'):
            byte_count = 3
            buffer += payload_byte
        elif(payload_byte[:5] == '11110'):
            byte_count = 4
            buffer += payload_byte
        elif(payload_byte[:2] == '10'):
            buffer += payload_byte
            continuation_byte_count += 1
            if(continuation_byte_count == byte_count - 1):
                decoded_payload += bytes(int(buffer[i:i+8], 2) for i in range(0, len(buffer), 8)).decode('utf-8')
                buffer = ""
                byte_count = 1
                continuation_byte_count = 0
        if decoded_payload[DELIMITER_IDX:] == DELIMITER:
            return decoded_payload[:DELIMITER_IDX]
    raise ValueError(
        "Could not decode. (Could not find the delimiter '%s')" % DELIMITER)


def encode_docx(text_path, output_path, secret_data, lsb=1):
    # Get contents of cover file
    doc = Document(text_path)

    # Get maximum bytes to encode (only encode before and after selected cover character)
    count_cover_char = 0
    for paragraph in doc.paragraphs:
        count_cover_char += paragraph.text.count(COVER_CHAR)
    max_encode_bytes = count_cover_char * 2 * lsb // 8
    print("[*] Maximum bytes to encode:", max_encode_bytes)

    # Append delimiter to payload
    secret_data += DELIMITER
    bytes_secret_data = bytes(secret_data, 'utf-8')

    # Check if payload is larger than maximum bytes
    if len(bytes_secret_data) > max_encode_bytes:
        raise ValueError(
            "[!] Insufficient bytes, need bigger docx file or less data.")

    # Convert payload to binary string
    print("[*] Encoding data...")
    binary_secret_data = ''.join(i for i in to_bin(bytes_secret_data))

    # Initialise sentinel values
    len_payload = len(binary_secret_data)
    idx_payload = 0
    base_steg_char = ''.join(i for i in to_bin(
        bytes(STEG_CHAR, 'utf-8')))[:-lsb]

    # Loop through each paragraph in document
    for paragraph in doc.paragraphs:
        # Get new runs
        list_runs = []
        for run in paragraph.runs:
            # Check if payload is already encoded and for cover character
            if (COVER_CHAR in run.text and idx_payload < len_payload):
                # Split text
                list_text = run.text.split(COVER_CHAR)
                # Loop through each text
                for text in list_text:
                    # Add run of cover text
                    list_runs.append(paragraph.add_run(text))

                    # Get next set of bits to encode (padded with 0 if not enough bits)
                    encode = binary_secret_data[idx_payload:idx_payload +
                                                lsb].ljust(lsb, '0')
                    # Get payload character
                    payload_char = bytes.fromhex(
                        hex(int(base_steg_char + encode, 2))[2:]).decode('utf-8')
                    # Add payload character to list of runs
                    before = paragraph.add_run(payload_char)
                    before.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    before.font.hidden = True
                    list_runs.append(before)
                    idx_payload += lsb

                    # Add cover character
                    list_runs.append(paragraph.add_run(COVER_CHAR))

                    # Check if end of payload is reached before encoding again (two encodings are performed for each cover character, before and after it)
                    if (idx_payload < len_payload):
                        # Get next set of bits to encode (padded with 0 if not enough bits)
                        encode = binary_secret_data[idx_payload:idx_payload + lsb].ljust(
                            lsb, '0')
                        # Get payload character
                        payload_char = bytes.fromhex(
                            hex(int(base_steg_char + encode, 2))[2:]).decode('utf-8')

                        # Add payload character to list of runs
                        after = paragraph.add_run(payload_char)
                        after.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        after.font.hidden = True
                        list_runs.append(after)
                        idx_payload += lsb
            else:
                list_runs.append(run)

        # Clear runs
        for run in paragraph.runs:
            paragraph._p.remove(run._r)

        # Add new runs
        for run in list_runs:
            paragraph._p.append(run._r)

    # Write stego text to output file
    doc.save(output_path)


def decode_docx(text_path, lsb=1):
    # Get contents of stego file
    doc = Document(text_path)
    steg_text = ""
    for paragraph in doc.paragraphs:
        steg_text += paragraph.text
    print("[+] Decoding...")

    # Initialise sentinel values
    esc_cover_char = re.escape(COVER_CHAR)
    regex = r"[^{}]{}[^{}]".format(
        esc_cover_char, esc_cover_char, esc_cover_char)
    base_steg_char = ''.join(i for i in to_bin(
        bytes(STEG_CHAR, 'utf-8')))[:-lsb]
    binary_payload = ""

    # Get payload characters
    list_payload_char = re.findall(regex, steg_text)
    for raw in list_payload_char:
        list_raw = raw.split(COVER_CHAR)
        for char in list_raw:
            # Convert char to binary
            char = ''.join(i for i in to_bin(bytes(char, 'utf-8')))
            # Populate list of steg characters
            if (char[:-lsb] == base_steg_char):
                binary_payload += char[-lsb:]

    # Convert payload to list of bytes
    list_payload_bytes = [binary_payload[i:i + 8]
                          for i in range(0, len(binary_payload), 8)]

    # Convert payload to characters
    decoded_payload = ""
    buffer = ""
    byte_count = 1
    continuation_byte_count = 0
    for payload_byte in list_payload_bytes:
        # Check variable byte encoding
        if(payload_byte[0] == '0'):
            byte_count = 1
            decoded_payload += chr(int(payload_byte, 2))
        elif(payload_byte[:3] == '110'):
            byte_count = 2
            buffer += payload_byte
        elif(payload_byte[:4] == '1110'):
            byte_count = 3
            buffer += payload_byte
        elif(payload_byte[:5] == '11110'):
            byte_count = 4
            buffer += payload_byte
        elif(payload_byte[:2] == '10'):
            buffer += payload_byte
            continuation_byte_count += 1
            if(continuation_byte_count == byte_count - 1):
                decoded_payload += bytes(int(buffer[i:i+8], 2) for i in range(0, len(buffer), 8)).decode('utf-8')
                buffer = ""
                byte_count = 1
                continuation_byte_count = 0
        if decoded_payload[DELIMITER_IDX:] == DELIMITER:
            return decoded_payload[:DELIMITER_IDX]
    raise ValueError(
        "Could not decode. (Could not find the delimiter '%s')" % DELIMITER)


def encode_xlsx(text_path, output_path, secret_data, lsb=1):
    # Load workbook and create hidden sheet
    wb = openpyxl.load_workbook(text_path)
    ws = wb.create_sheet(DELIMITER)
    ws.sheet_state = 'hidden'

    # Convert payload to binary string
    print("[*] Encoding data...")
    # Append delimiter to payload
    secret_data += DELIMITER
    bytes_secret_data = bytes(secret_data, 'utf-8')
    binary_secret_data = ''.join(i for i in to_bin(bytes_secret_data))

    # Initialise sentinel values
    steg_text = ''
    base_steg_char = ''.join(i for i in to_bin(
        bytes(STEG_CHAR, 'utf-8')))[:-lsb]

    # Loop through each bit of secret data
    for idx_payload in range(0, len(binary_secret_data), lsb):
        # Get next set of bits to encode (padded with 0 if not enough bits)
        encode = binary_secret_data[idx_payload:idx_payload +
                                    lsb].ljust(lsb, '0')
        # Get payload character
        payload_char = bytes.fromhex(
            hex(int(base_steg_char + encode, 2))[2:]).decode('utf-8')
        # Add payload character to stego text
        steg_text += payload_char

    # Store payload in hidden worksheet
    idx_payload = len(DELIMITER)
    ws.cell(row=idx_payload, column=idx_payload).value = steg_text
    ws.cell(row=idx_payload, column=idx_payload).font = Font(color="FFFFFF")

    # Save workbook
    wb.save(output_path)


def decode_xlsx(text_path, lsb=1):
    # Get contents of stego file
    wb = openpyxl.load_workbook(text_path)
    ws = wb[DELIMITER]
    idx_payload = len(DELIMITER)
    steg_text = ws.cell(row=idx_payload, column=idx_payload).value
    print("[+] Decoding...")

    # Initialise sentinel values
    base_steg_char = ''.join(i for i in to_bin(
        bytes(STEG_CHAR, 'utf-8')))[:-lsb]
    binary_payload = ""

    # Get payload characters
    for char in steg_text:
        # Convert char to binary
        char = ''.join(i for i in to_bin(bytes(char, 'utf-8')))
        # Populate list of steg characters
        if (char[:-lsb] == base_steg_char):
            binary_payload += char[-lsb:]

    # Convert payload to list of bytes
    list_payload_bytes = [binary_payload[i:i + 8]
                          for i in range(0, len(binary_payload), 8)]

    # Convert payload to characters
    decoded_payload = ""
    buffer = ""
    byte_count = 1
    continuation_byte_count = 0
    for payload_byte in list_payload_bytes:
        # Check variable byte encoding
        if(payload_byte[0] == '0'):
            byte_count = 1
            decoded_payload += chr(int(payload_byte, 2))
        elif(payload_byte[:3] == '110'):
            byte_count = 2
            buffer += payload_byte
        elif(payload_byte[:4] == '1110'):
            byte_count = 3
            buffer += payload_byte
        elif(payload_byte[:5] == '11110'):
            byte_count = 4
            buffer += payload_byte
        elif(payload_byte[:2] == '10'):
            buffer += payload_byte
            continuation_byte_count += 1
            if(continuation_byte_count == byte_count - 1):
                decoded_payload += bytes(int(buffer[i:i+8], 2) for i in range(0, len(buffer), 8)).decode('utf-8')
                buffer = ""
                byte_count = 1
                continuation_byte_count = 0
        if decoded_payload[DELIMITER_IDX:] == DELIMITER:
            return decoded_payload[:DELIMITER_IDX]
    raise ValueError(
        "Could not decode. (Could not find the delimiter '%s')" % DELIMITER)


def encode_pdf(text_path, output_path, secret_data, lsb=1):
    # Open cover file
    pdf = PdfReader(text_path)

    # Convert payload to binary string
    print("[*] Encoding data...")
    # Append delimiter to payload
    secret_data += DELIMITER
    bytes_secret_data = bytes(secret_data, 'utf-8')
    binary_secret_data = ''.join(i for i in to_bin(bytes_secret_data))

    # Initialise sentinel values
    steg_text = ''
    base_steg_char = ''.join(i for i in to_bin(
        bytes(STEG_CHAR, 'utf-8')))[:-lsb]

    # Loop through each bit of secret data
    for idx_payload in range(0, len(binary_secret_data), lsb):
        # Get next set of bits to encode (padded with 0 if not enough bits)
        encode = binary_secret_data[idx_payload:idx_payload +
                                    lsb].ljust(lsb, '0')
        # Get payload character
        payload_char = bytes.fromhex(
            hex(int(base_steg_char + encode, 2))[2:]).decode('utf-8')
        # Add payload character to stego text
        steg_text += payload_char

    # Initialise and update output pdf
    output_pdf = PdfWriter()
    output_pdf.append_pages_from_reader(pdf)
    output_pdf.add_metadata({**pdf.metadata, '/' + DELIMITER: steg_text})

    # Save output pdf
    output_pdf.write(output_path)


def decode_pdf(text_path, lsb=1):
    # Get contents of stego file
    pdf = PdfReader(text_path)

    # Check for hidden metadata
    if (('/' + DELIMITER) in pdf.metadata.keys()):
        steg_text = pdf.metadata['/' + DELIMITER]
        print("[+] Decoding...")

        # Initialise sentinel values
        base_steg_char = ''.join(i for i in to_bin(
            bytes(STEG_CHAR, 'utf-8')))[:-lsb]
        binary_payload = ""

        # Get payload characters
        for char in steg_text:
            # Convert char to binary
            char = ''.join(i for i in to_bin(bytes(char, 'utf-8')))
            # Populate list of steg characters
            if (char[:-lsb] == base_steg_char):
                binary_payload += char[-lsb:]

        # Convert payload to list of bytes
        list_payload_bytes = [binary_payload[i:i + 8]
                              for i in range(0, len(binary_payload), 8)]

        # Convert payload to characters
        decoded_payload = ""
        buffer = ""
        byte_count = 1
        continuation_byte_count = 0
        for payload_byte in list_payload_bytes:
            # Check variable byte encoding
            if(payload_byte[0] == '0'):
                byte_count = 1
                decoded_payload += chr(int(payload_byte, 2))
            elif(payload_byte[:3] == '110'):
                byte_count = 2
                buffer += payload_byte
            elif(payload_byte[:4] == '1110'):
                byte_count = 3
                buffer += payload_byte
            elif(payload_byte[:5] == '11110'):
                byte_count = 4
                buffer += payload_byte
            elif(payload_byte[:2] == '10'):
                buffer += payload_byte
                continuation_byte_count += 1
                if(continuation_byte_count == byte_count - 1):
                    decoded_payload += bytes(int(buffer[i:i+8], 2) for i in range(0, len(buffer), 8)).decode('utf-8')
                    buffer = ""
                    byte_count = 1
                    continuation_byte_count = 0
            if decoded_payload[DELIMITER_IDX:] == DELIMITER:
                return decoded_payload[:DELIMITER_IDX]
        raise ValueError(
            "Could not decode. (Could not find the delimiter '%s')" % DELIMITER)
    else:
        return ValueError("Could not decode. (Could not find any hidden data)")


def encode_mp4(mp4_path, output_path, secret_data, lsb=1):
    # Convert payload to binary string
    print("[*] Encoding data...")
    # Append delimiter to payload
    secret_data += DELIMITER
    bytes_secret_data = bytes(secret_data, 'utf-8')
    binary_secret_data = ''.join(i for i in to_bin(bytes_secret_data))

    # Initialise sentinel values
    steg_text = ''
    base_steg_char = ''.join(i for i in to_bin(
        bytes(STEG_CHAR, 'utf-8')))[:-lsb]

    # Loop through each bit of secret data
    for idx_payload in range(0, len(binary_secret_data), lsb):
        # Get next set of bits to encode (padded with 0 if not enough bits)
        encode = binary_secret_data[idx_payload:idx_payload +
                                    lsb].ljust(lsb, '0')
        # Get payload character
        payload_char = bytes.fromhex(
            hex(int(base_steg_char + encode, 2))[2:]).decode('utf-8')
        # Add payload character to stego text
        steg_text += payload_char

    ffmpeg.input(mp4_path).output(
        output_path, metadata=f"MediaProperties={steg_text}", movflags='use_metadata_tags', loglevel='quiet').overwrite_output().run()


def decode_mp4(mp4_path, lsb=1):
    # Check for hidden metadata
    if "MediaProperties" in ffmpeg.probe(mp4_path)['format']['tags'].keys():
        steg_text = ffmpeg.probe(mp4_path)['format']['tags']['MediaProperties']
        print("[+] Decoding...")

        # Initialise sentinel values
        base_steg_char = ''.join(i for i in to_bin(
            bytes(STEG_CHAR, 'utf-8')))[:-lsb]
        binary_payload = ""

        # Get payload characters
        for char in steg_text:
            # Convert char to binary
            char = ''.join(i for i in to_bin(bytes(char, 'utf-8')))
            # Populate list of steg characters
            if (char[:-lsb] == base_steg_char):
                binary_payload += char[-lsb:]

        # Convert payload to list of bytes
        list_payload_bytes = [binary_payload[i:i + 8]
                              for i in range(0, len(binary_payload), 8)]

        # Convert payload to characters
        decoded_payload = ""
        buffer = ""
        byte_count = 1
        continuation_byte_count = 0
        for payload_byte in list_payload_bytes:
            # Check variable byte encoding
            if(payload_byte[0] == '0'):
                byte_count = 1
                decoded_payload += chr(int(payload_byte, 2))
            elif(payload_byte[:3] == '110'):
                byte_count = 2
                buffer += payload_byte
            elif(payload_byte[:4] == '1110'):
                byte_count = 3
                buffer += payload_byte
            elif(payload_byte[:5] == '11110'):
                byte_count = 4
                buffer += payload_byte
            elif(payload_byte[:2] == '10'):
                buffer += payload_byte
                continuation_byte_count += 1
                if(continuation_byte_count == byte_count - 1):
                    decoded_payload += bytes(int(buffer[i:i+8], 2) for i in range(0, len(buffer), 8)).decode('utf-8')
                    buffer = ""
                    byte_count = 1
                    continuation_byte_count = 0
            if decoded_payload[DELIMITER_IDX:] == DELIMITER:
                return decoded_payload[:DELIMITER_IDX]
        raise ValueError(
            f"Could not decode. (Could not find the delimiter {DELIMITER})")
    else:
        raise ValueError("Could not decode. (Could not find any hidden data)")


def encode(delimeter, filepath, filename, text, lsb=1):
    global DELIMITER
    global DELIMITER_IDX

    DELIMITER = delimeter
    DELIMITER_IDX = -len(DELIMITER)
    input = os.path.join(filepath, filename)
    output = os.path.join(filepath, "enc_" + filename)

    # Image encoding
    if input.endswith(".png") or input.endswith(".bmp"):
        # encode the data into the image
        encoded_image = encode_image(input, text, lsb)
        # save the output image (encoded image)
        cv2.imwrite(output, encoded_image)
        print("[+] Saved encoded image.")
    # Audio encoding
    elif input.endswith(".wav"):
        # encode the data into the audio
        encode_audio(input, output, text, lsb)
        print("[+] Saved encoded audio.")
    # Text encoding
    elif input.endswith(".txt"):
        encode_text(input, output, text, lsb)
        print("[+] Saved encoded text.")
    # GIF encoding
    elif input.endswith(".gif"):
        encode_gif(input, output, text, lsb)
        print("[+] Saved encoded GIF.")
    # DOCX encoding
    elif input.endswith(".docx"):
        encode_docx(input, output, text, lsb)
        print("[+] Saved encoded document.")
    # XLSX encoding
    elif input.endswith(".xlsx"):
        encode_xlsx(input, output, text, lsb)
        print("[+] Saved encoded document.")
    # PDF encoding
    elif input.endswith(".pdf"):
        encode_pdf(input, output, text, lsb)
        print("[+] Saved encoded document.")
    # MP4 encoding
    elif input.endswith(".mp4"):
        encode_mp4(input, output, text, lsb)
        print("[+] Saved encoded video.")
    else:
        print("File type not supported!")


def decode(delimiter, filepath, filename, lsb):
    global DELIMITER
    global DELIMITER_IDX

    DELIMITER = delimiter
    DELIMITER_IDX = -len(DELIMITER)
    input = os.path.join(filepath, filename)

    if input.endswith(".png") or input.endswith(".bmp"):
        decoded_data = decode_image(input, lsb)
    elif input.endswith(".wav"):
        decoded_data = decode_audio(input, lsb)
    elif input.endswith(".txt"):
        decoded_data = decode_text(input, lsb)
    elif input.endswith(".gif"):
        decoded_data = decode_gif(input, lsb)
    elif input.endswith(".docx"):
        decoded_data = decode_docx(input, lsb)
    elif input.endswith(".xlsx"):
        decoded_data = decode_xlsx(input, lsb)
    elif input.endswith(".pdf"):
        decoded_data = decode_pdf(input, lsb)
    elif input.endswith(".mp4"):
        decoded_data = decode_mp4(input, lsb)

    # Output results
    if decoded_data:
        print("[+] Decoded data:", decoded_data)
        return decoded_data
    else:
        print("[!] Could not decode.")


def html_display(filename):
    if filename.endswith(('.png', '.bmp', '.gif')):
        return f'<img src="../static/temp/{filename}">'
    elif filename.endswith('.wav'):
        return f'<audio src="../static/temp/{filename}" controls></audio>'
    elif filename.endswith(".txt"):
        return f'<iframe src="../static/temp/{filename}"></iframe>'
    elif filename.endswith(".docx"):
        return f'<i class="fa-solid fa-file-word"></i>'
    elif filename.endswith(".xlsx"):
        return f'<i class="fa-solid fa-file-excel"></i>'
    elif filename.endswith(".pdf"):
        return f'<embed src="../static/temp/{filename}" type="application/pdf">'
    elif filename.endswith(".mp4"):
        return f'<video controls><source src="../static/temp/{filename}" type="video/mp4"></video>'
    else:
        return 'File type not supported'
